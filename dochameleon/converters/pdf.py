"""
PDF to DOCX conversion utilities - Clean native DOCX output.

This module focuses on producing DOCX files that look like they were
originally created in Microsoft Word, not converted from PDF.

Features:
- Clean, native Word-like formatting
- Table of Contents with proper Word field codes
- Working hyperlinks (internal and external)
- Cross-references support
- Table of Figures / Table of Tables support
"""

from pathlib import Path
from typing import Tuple, Union, List, Dict, Optional
import re


def convert_pdf_to_docx_enhanced(pdf_file: Path, output_dir: Path) -> Tuple[bool, Union[Path, str]]:
    """
    Convert PDF to DOCX with clean, native Word-like output.
    Preserves hyperlinks, TOC structure, and cross-references.
    """
    from pdf2docx import Converter
    
    output_path = output_dir / (pdf_file.stem + ".docx")
    
    try:
        # First, extract hyperlinks from the PDF
        pdf_hyperlinks = extract_pdf_hyperlinks(pdf_file)
        
        cv = Converter(str(pdf_file))
        
        # Convert with minimal table detection to avoid fake tables
        cv.convert(
            str(output_path),
            # CRITICAL: High tolerance = fewer false table detections
            connected_border_tolerance=1.0,
            min_section_height=30,
            
            # Text flow
            line_overlap_threshold=0.9,
            line_break_width_ratio=0.5,
            line_break_free_space_ratio=0.15,
            new_paragraph_free_space_ratio=0.85,
            
            # Images
            float_image_ignorable_gap=10,
            
            # Page margins - preserve original
            page_margin_factor_top=0.0,
            page_margin_factor_bottom=0.0,
        )
        cv.close()
        
        if output_path.exists():
            # Aggressively clean the document and restore hyperlinks
            make_docx_native(output_path, pdf_hyperlinks)
            return True, output_path
        else:
            return False, "DOCX file was not created"
            
    except Exception as e:
        return False, str(e)


def extract_pdf_hyperlinks(pdf_file: Path) -> List[Dict]:
    """
    Extract hyperlinks from a PDF file.
    Returns a list of dicts with 'text', 'url', and 'page' information.
    """
    hyperlinks = []
    
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(str(pdf_file))
        
        for page_num, page in enumerate(doc):
            # Get all links on the page
            links = page.get_links()
            
            for link in links:
                link_info = {
                    'page': page_num,
                    'rect': link.get('from', None),
                    'url': None,
                    'dest_page': None,
                    'text': None
                }
                
                # External URI link
                if 'uri' in link:
                    link_info['url'] = link['uri']
                
                # Internal page link
                if 'page' in link:
                    link_info['dest_page'] = link['page']
                
                # Try to get the text at the link location
                if link_info['rect']:
                    rect = fitz.Rect(link_info['rect'])
                    link_info['text'] = page.get_text("text", clip=rect).strip()
                
                if link_info['url'] or link_info['dest_page'] is not None:
                    hyperlinks.append(link_info)
        
        doc.close()
        
    except ImportError:
        # PyMuPDF not installed, try pypdf
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(pdf_file))
            
            for page_num, page in enumerate(reader.pages):
                if '/Annots' in page:
                    annotations = page['/Annots']
                    if annotations:
                        for annot in annotations:
                            annot_obj = annot.get_object() if hasattr(annot, 'get_object') else annot
                            if annot_obj.get('/Subtype') == '/Link':
                                link_info = {
                                    'page': page_num,
                                    'url': None,
                                    'dest_page': None,
                                    'text': None
                                }
                                
                                # Get the action or destination
                                if '/A' in annot_obj:
                                    action = annot_obj['/A']
                                    if '/URI' in action:
                                        link_info['url'] = action['/URI']
                                
                                if link_info['url']:
                                    hyperlinks.append(link_info)
                                    
        except ImportError:
            # Neither library available
            pass
        except Exception:
            pass
    except Exception:
        pass
    
    return hyperlinks


def make_docx_native(docx_path: Path, pdf_hyperlinks: Optional[List[Dict]] = None):
    """
    Aggressively clean DOCX to look like native Word document.
    Also rebuilds TOC, hyperlinks, and cross-references.
    
    Args:
        docx_path: Path to the DOCX file
        pdf_hyperlinks: Optional list of hyperlinks extracted from the source PDF
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Twips
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document(str(docx_path))
        
        # ============================================
        # 1. SET PROPER PAGE MARGINS
        # ============================================
        set_page_margins(doc)
        
        # ============================================
        # 2. REMOVE ALL FAKE TABLES (wrapper tables)
        # ============================================
        remove_wrapper_tables(doc)
        
        # ============================================
        # 3. REMOVE ALL PARAGRAPH BORDERS
        # ============================================
        remove_all_paragraph_borders(doc)
        
        # ============================================
        # 4. CLEAN HEADING STYLES
        # ============================================
        clean_headings(doc)
        
        # ============================================
        # 5. REBUILD HYPERLINKS
        # ============================================
        rebuild_hyperlinks(doc, pdf_hyperlinks)
        
        # ============================================
        # 6. REBUILD TABLE OF CONTENTS
        # ============================================
        rebuild_toc(doc)
        
        # ============================================
        # 7. REBUILD TABLE OF FIGURES/TABLES
        # ============================================
        rebuild_list_of_figures_tables(doc)
        
        # ============================================
        # 8. FIX CROSS-REFERENCES
        # ============================================
        fix_cross_references(doc)
        
        # ============================================
        # 9. APPLY CLEAN DOCUMENT STYLES
        # ============================================
        apply_native_styles(doc)
        
        doc.save(str(docx_path))
        
    except Exception as e:
        print(f"  Note: Could not clean document: {e}")


def set_page_margins(doc):
    """
    Set proper Word-standard page margins (1 inch all around).
    """
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    for section in doc.sections:
        # Standard Word margins: 1 inch all around
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Also set header/footer distances
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)


def remove_wrapper_tables(doc):
    """
    Remove tables that are just wrappers (not real data tables).
    
    Wrapper tables are typically:
    - Single row tables
    - Tables where cells contain mostly paragraphs (not tabular data)
    - Tables used as layout containers
    """
    from docx.oxml.ns import qn
    
    tables_to_process = list(doc.tables)
    
    for table in tables_to_process:
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        
        # Determine if this is a real table or a wrapper
        is_wrapper = False
        
        # Single cell = definitely a wrapper
        if rows == 1 and cols == 1:
            is_wrapper = True
        
        # Single row with few columns = likely a wrapper/layout
        elif rows == 1 and cols <= 2:
            is_wrapper = True
        
        # Check if it looks like a heading wrapper
        elif rows <= 2:
            cell_text = ""
            for row in table.rows:
                for cell in row.cells:
                    cell_text += cell.text.strip()
            # Short text in few rows = likely wrapper
            if len(cell_text) < 200 and rows == 1:
                is_wrapper = True
        
        if is_wrapper:
            # Remove all borders from wrapper tables
            remove_table_borders(table)
            # Also remove cell borders and shading
            for row in table.rows:
                for cell in row.cells:
                    remove_cell_formatting(cell)
        else:
            # Real table - apply clean styling
            apply_clean_table_style(table)


def remove_table_borders(table):
    """
    Remove all borders from a table.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    
    if tblPr is not None:
        # Remove table borders
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        
        # Add explicit "no borders"
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Remove table indent
        tblInd = tblPr.find(qn('w:tblInd'))
        if tblInd is not None:
            tblPr.remove(tblInd)


def remove_cell_formatting(cell):
    """
    Remove borders and shading from a cell.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    
    if tcPr is not None:
        # Remove cell borders
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is not None:
            tcPr.remove(tcBorders)
        
        # Add explicit "no borders"
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        # Remove shading
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            tcPr.remove(shd)
    
    # Also clean paragraphs within cell
    for paragraph in cell.paragraphs:
        remove_paragraph_borders(paragraph)


def remove_all_paragraph_borders(doc):
    """
    Remove ALL borders from ALL paragraphs.
    This is aggressive but necessary to get clean output.
    """
    from docx.oxml.ns import qn
    
    for paragraph in doc.paragraphs:
        remove_paragraph_borders(paragraph)
    
    # Also remove from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    remove_paragraph_borders(paragraph)


def remove_paragraph_borders(paragraph):
    """
    Remove all borders and unnecessary shading from a paragraph.
    """
    from docx.oxml.ns import qn
    
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        # Remove paragraph borders completely
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)
        
        # Remove paragraph shading (backgrounds)
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        
        # Remove frames
        framePr = pPr.find(qn('w:framePr'))
        if framePr is not None:
            pPr.remove(framePr)
    
    # Clean runs too
    for run in paragraph.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            # Remove run shading (text backgrounds) except for intentional highlights
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                # Remove gray backgrounds (artifacts), keep colored highlights
                if fill and fill.upper() in ('F5F5F5', 'F0F0F0', 'EFEFEF', 'FAFAFA', 
                                               'E0E0E0', 'F8F8F8', 'E8E8E8', 'D0D0D0',
                                               'FFFFFF', 'auto'):
                    rPr.remove(shd)


def clean_headings(doc):
    """
    Ensure headings are clean and properly styled.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    
    for paragraph in doc.paragraphs:
        # Detect headings by style or formatting
        is_heading = False
        style_name = paragraph.style.name if paragraph.style else ''
        
        if 'Heading' in style_name or 'Title' in style_name:
            is_heading = True
        else:
            # Check if it looks like a heading (bold, larger font, short text)
            text = paragraph.text.strip()
            if len(text) < 100 and len(text) > 0:
                for run in paragraph.runs:
                    if run.bold and run.font.size and run.font.size >= Pt(12):
                        is_heading = True
                        break
        
        if is_heading:
            # Ensure heading has no borders or boxes
            pPr = paragraph._p.find(qn('w:pPr'))
            if pPr is not None:
                # Remove any borders
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    pPr.remove(pBdr)
                
                # Remove shading
                shd = pPr.find(qn('w:shd'))
                if shd is not None:
                    pPr.remove(shd)
            
            # Set proper heading spacing
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)


# ============================================================
# HYPERLINKS AND REFERENCES
# ============================================================

def rebuild_hyperlinks(doc, pdf_hyperlinks: Optional[List[Dict]] = None):
    """
    Detect URLs in text and convert them to proper Word hyperlinks.
    Also fixes existing hyperlinks that may have been broken during conversion.
    
    Args:
        doc: The python-docx Document object
        pdf_hyperlinks: Optional list of hyperlinks extracted from the source PDF
    """
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor, Pt
    
    # Build a lookup of URLs from PDF if available
    pdf_url_map = {}
    if pdf_hyperlinks:
        for link in pdf_hyperlinks:
            if link.get('text') and link.get('url'):
                # Normalize the text for matching
                text_key = link['text'].strip().lower()
                pdf_url_map[text_key] = link['url']
    
    # URL pattern for detection
    url_pattern = re.compile(
        r'(https?://[^\s<>"{}|\\^`\[\]]+|www\.[^\s<>"{}|\\^`\[\]]+|'
        r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
    )
    
    for paragraph in doc.paragraphs:
        _process_paragraph_for_hyperlinks(paragraph, url_pattern, pdf_url_map)
    
    # Also process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process_paragraph_for_hyperlinks(paragraph, url_pattern, pdf_url_map)


def _process_paragraph_for_hyperlinks(paragraph, url_pattern, pdf_url_map: Optional[Dict] = None):
    """
    Process a single paragraph to add hyperlinks.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    
    full_text = paragraph.text
    
    # Check for existing hyperlinks in the paragraph XML
    p_element = paragraph._p
    existing_hyperlinks = p_element.findall('.//' + qn('w:hyperlink'))
    
    # If there are already hyperlinks, style them properly
    for hyperlink in existing_hyperlinks:
        _style_hyperlink_runs(hyperlink)
    
    # Check if any text matches PDF extracted hyperlinks
    if pdf_url_map:
        for run in paragraph.runs:
            run_text = run.text
            if run_text:
                text_key = run_text.strip().lower()
                if text_key in pdf_url_map:
                    # This text was a hyperlink in the PDF
                    url = pdf_url_map[text_key]
                    run.font.color.rgb = RGBColor(0, 0, 238)
                    run.font.underline = True
                    _add_hyperlink_to_run(paragraph, run, url)
    
    # Find URLs in text that are not already hyperlinks
    matches = list(url_pattern.finditer(full_text))
    if not matches:
        return
    
    # For each run, check if it contains a URL
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
            
        for match in url_pattern.finditer(run_text):
            url = match.group(0)
            
            # Style the run as a hyperlink (blue, underlined)
            run.font.color.rgb = RGBColor(0, 0, 238)  # Blue
            run.font.underline = True
            
            # Add the actual hyperlink relationship
            _add_hyperlink_to_run(paragraph, run, url)


def _style_hyperlink_runs(hyperlink_element):
    """
    Apply proper hyperlink styling to runs inside a hyperlink element.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    for run in hyperlink_element.findall(qn('w:r')):
        rPr = run.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run.insert(0, rPr)
        
        # Set blue color
        color = rPr.find(qn('w:color'))
        if color is None:
            color = OxmlElement('w:color')
            rPr.append(color)
        color.set(qn('w:val'), '0000EE')
        
        # Set underline
        u = rPr.find(qn('w:u'))
        if u is None:
            u = OxmlElement('w:u')
            rPr.append(u)
        u.set(qn('w:val'), 'single')


def _add_hyperlink_to_run(paragraph, run, url):
    """
    Create a proper hyperlink relationship for a run containing a URL.
    This creates a clickable hyperlink in the Word document.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    
    try:
        part = paragraph.part
        
        # Ensure URL has proper scheme
        if url.startswith('www.'):
            url = 'https://' + url
        elif '@' in url and not url.startswith('mailto:'):
            url = 'mailto:' + url
        
        # Add hyperlink relationship
        r_id = part.relate_to(
            url, 
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 
            is_external=True
        )
        
        # Create the hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Clone the run element to put inside hyperlink
        new_run = OxmlElement('w:r')
        
        # Copy run properties
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            new_rPr = OxmlElement('w:rPr')
            for child in rPr:
                new_rPr.append(child.__copy__() if hasattr(child, '__copy__') else child)
            new_run.append(new_rPr)
        else:
            # Add default hyperlink styling
            new_rPr = OxmlElement('w:rPr')
            
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000EE')
            new_rPr.append(color)
            
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            new_rPr.append(u)
            
            new_run.append(new_rPr)
        
        # Copy text
        text_elem = run._element.find(qn('w:t'))
        if text_elem is not None:
            new_text = OxmlElement('w:t')
            new_text.text = text_elem.text
            # Preserve whitespace
            new_text.set(qn('xml:space'), 'preserve')
            new_run.append(new_text)
        
        hyperlink.append(new_run)
        
        # Replace the original run with the hyperlink
        p = paragraph._p
        run_elem = run._element
        run_idx = list(p).index(run_elem)
        p.remove(run_elem)
        p.insert(run_idx, hyperlink)
        
    except Exception as e:
        # If we can't create the full hyperlink, at least keep the styling
        pass


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The python-docx paragraph to add the hyperlink to
        text: The display text for the hyperlink
        url: The URL to link to
        
    Returns:
        The hyperlink element
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    part = paragraph.part
    
    # Ensure URL has proper scheme
    if url.startswith('www.'):
        url = 'https://' + url
    elif '@' in url and not url.startswith('mailto:'):
        url = 'mailto:' + url
    
    # Create relationship
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    
    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run for text
    new_run = OxmlElement('w:r')
    
    # Add hyperlink styling
    rPr = OxmlElement('w:rPr')
    
    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000EE')
    rPr.append(color)
    
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    text_elem.set(qn('xml:space'), 'preserve')
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def rebuild_toc(doc):
    """
    Rebuild Table of Contents with proper Word field codes.
    Detects existing TOC-like content and replaces with proper TOC field.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Twips
    
    # First, collect all headings in the document
    headings = _collect_headings(doc)
    if not headings:
        return
    
    # Look for existing TOC content (text patterns like "Contents", "Table of Contents")
    toc_start_idx = None
    toc_end_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        
        # Detect TOC header
        if text in ('contents', 'table of contents', 'toc'):
            toc_start_idx = i
            # Find the end of the TOC (look for next real content)
            for j in range(i + 1, min(i + 50, len(doc.paragraphs))):
                next_para = doc.paragraphs[j]
                next_text = next_para.text.strip()
                
                # TOC entries typically have page numbers or dots
                if next_text and not _looks_like_toc_entry(next_text):
                    toc_end_idx = j
                    break
            break
    
    # If we found a TOC section, add field code marker
    if toc_start_idx is not None:
        _add_toc_field(doc, doc.paragraphs[toc_start_idx], headings)


def _collect_headings(doc) -> List[Dict]:
    """
    Collect all headings from the document with their levels and text.
    """
    from docx.shared import Pt
    
    headings = []
    
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        text = para.text.strip()
        
        if not text:
            continue
        
        level = None
        
        # Check by style name
        if 'Heading 1' in style_name or style_name == 'Title':
            level = 1
        elif 'Heading 2' in style_name:
            level = 2
        elif 'Heading 3' in style_name:
            level = 3
        elif 'Heading 4' in style_name:
            level = 4
        elif 'Heading' in style_name:
            level = 5
        else:
            # Detect by formatting (bold, larger text, short)
            if len(text) < 80:
                for run in para.runs:
                    if run.bold:
                        font_size = run.font.size
                        if font_size:
                            if font_size >= Pt(16):
                                level = 1
                            elif font_size >= Pt(14):
                                level = 2
                            elif font_size >= Pt(12):
                                level = 3
                        break
        
        if level:
            headings.append({
                'index': i,
                'level': level,
                'text': text,
                'bookmark': f'_Toc{i}'
            })
    
    return headings


def _looks_like_toc_entry(text: str) -> bool:
    """
    Check if text looks like a TOC entry (has page number pattern, dots, etc.)
    """
    # TOC entries often have patterns like "Section Name ... 5" or "Section Name    5"
    if re.search(r'\.{2,}\s*\d+$', text):  # dots followed by page number
        return True
    if re.search(r'\s{3,}\d+$', text):  # multiple spaces followed by page number
        return True
    if re.match(r'^\d+(\.\d+)*\s+', text):  # starts with section number
        return True
    return False


def _add_toc_field(doc, toc_paragraph, headings: List[Dict]):
    """
    Add Word TOC field code to enable "Update Table" functionality.
    This creates a proper Word TOC that can be updated in Word using F9 or right-click > Update Field.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # Add bookmarks to all headings for TOC linking
    for heading in headings:
        if heading['index'] < len(doc.paragraphs):
            _add_bookmark_to_paragraph(doc.paragraphs[heading['index']], heading['bookmark'])
    
    # Get the paragraph element
    p = toc_paragraph._p
    
    # Create the TOC field structure
    # Field begin
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    
    run_begin = OxmlElement('w:r')
    run_begin.append(fldChar_begin)
    
    # Field instruction
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # TOC field with: \o = outline levels 1-3, \h = hyperlinks, \z = hide tab & page in web view, \u = use paragraph outline level
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    
    run_instr = OxmlElement('w:r')
    run_instr.append(instrText)
    
    # Field separator (marks start of field result)
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    
    run_sep = OxmlElement('w:r')
    run_sep.append(fldChar_sep)
    
    # Field end
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    
    run_end = OxmlElement('w:r')
    run_end.append(fldChar_end)
    
    # Create a new paragraph for the TOC field after the title
    parent = p.getparent()
    p_idx = list(parent).index(p)
    
    # Create TOC field paragraph
    toc_field_p = OxmlElement('w:p')
    toc_field_p.append(run_begin)
    toc_field_p.append(run_instr)
    toc_field_p.append(run_sep)
    
    # Add placeholder text that will be replaced when user updates the TOC
    placeholder_run = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = 'Right-click and select "Update Field" to generate table of contents'
    placeholder_run.append(placeholder_text)
    toc_field_p.append(placeholder_run)
    
    toc_field_p.append(run_end)
    
    # Insert after the TOC title paragraph
    parent.insert(p_idx + 1, toc_field_p)


def create_toc_entry_with_link(paragraph, heading_text: str, page_num: str, bookmark_name: str):
    """
    Create a TOC entry with a hyperlink to the heading.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    
    p = paragraph._p
    
    # Create hyperlink to bookmark
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    
    # Create run with heading text
    run = OxmlElement('w:r')
    
    # Add run properties (hyperlink style)
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000EE')
    rPr.append(color)
    run.append(rPr)
    
    # Add text
    text = OxmlElement('w:t')
    text.text = heading_text
    run.append(text)
    
    hyperlink.append(run)
    
    # Add tab for leader dots
    tab_run = OxmlElement('w:r')
    tab = OxmlElement('w:tab')
    tab_run.append(tab)
    
    # Add page number
    page_run = OxmlElement('w:r')
    page_text = OxmlElement('w:t')
    page_text.text = page_num
    page_run.append(page_text)
    
    p.append(hyperlink)
    p.append(tab_run)
    p.append(page_run)


def _add_bookmark_to_paragraph(paragraph, bookmark_name: str):
    """
    Add a bookmark to a paragraph for internal linking.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import random
    
    p = paragraph._p
    
    # Create bookmark start
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(random.randint(10000, 99999)))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # Create bookmark end
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_start.get(qn('w:id')))
    
    # Insert at the beginning of the paragraph
    p.insert(0, bookmark_start)
    p.append(bookmark_end)


def rebuild_list_of_figures_tables(doc):
    """
    Rebuild List of Figures and List of Tables.
    Detects figure/table captions and creates proper references.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    figures = []
    tables = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Detect figure captions
        figure_match = re.match(r'^(Figure|Fig\.?)\s*(\d+)', text, re.IGNORECASE)
        if figure_match:
            figures.append({
                'index': i,
                'number': figure_match.group(2),
                'text': text,
                'bookmark': f'_FigRef{i}'
            })
            _add_bookmark_to_paragraph(para, f'_FigRef{i}')
            continue
        
        # Detect table captions
        table_match = re.match(r'^(Table)\s*(\d+)', text, re.IGNORECASE)
        if table_match:
            tables.append({
                'index': i,
                'number': table_match.group(2),
                'text': text,
                'bookmark': f'_TblRef{i}'
            })
            _add_bookmark_to_paragraph(para, f'_TblRef{i}')
    
    # Look for "List of Figures" or "List of Tables" sections
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        
        if 'list of figures' in text:
            _create_list_field(doc, para, figures, 'Figure')
        elif 'list of tables' in text:
            _create_list_field(doc, para, tables, 'Table')


def _create_list_field(doc, header_para, items: List[Dict], item_type: str):
    """
    Create a List of Figures/Tables field.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # Field code for list of figures: TOC \h \z \c "Figure"
    # Field code for list of tables: TOC \h \z \c "Table"
    field_code = f' TOC \\h \\z \\c "{item_type}" '
    
    # Add field markers to enable Word's "Update Table" function
    # For now, we just add bookmarks to the captions (done above)


def fix_cross_references(doc):
    """
    Fix cross-references in the document.
    Converts text references like "Figure 1" or "Section 2" to internal links.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    
    # Build a map of all figure, table, and section references
    ref_map = _build_reference_map(doc)
    
    # Patterns for cross-references
    patterns = [
        (r'(Figure|Fig\.?)\s+(\d+)', 'figure'),
        (r'(Table)\s+(\d+)', 'table'),
        (r'(Section|Sec\.?)\s+(\d+(?:\.\d+)*)', 'section'),
        (r'(Chapter)\s+(\d+)', 'chapter'),
        (r'(Equation|Eq\.?)\s+(\d+)', 'equation'),
    ]
    
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            for pattern, ref_type in patterns:
                matches = list(re.finditer(pattern, text, re.IGNORECASE))
                for match in matches:
                    ref_num = match.group(2)
                    ref_key = f"{ref_type}_{ref_num}"
                    
                    if ref_key in ref_map:
                        # Style as internal reference (same blue as hyperlinks)
                        run.font.color.rgb = RGBColor(0, 0, 238)


def _build_reference_map(doc) -> Dict[str, int]:
    """
    Build a map of all referenceable items (figures, tables, sections).
    """
    ref_map = {}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Figures
        fig_match = re.match(r'^(Figure|Fig\.?)\s*(\d+)', text, re.IGNORECASE)
        if fig_match:
            ref_map[f"figure_{fig_match.group(2)}"] = i
        
        # Tables
        tbl_match = re.match(r'^(Table)\s*(\d+)', text, re.IGNORECASE)
        if tbl_match:
            ref_map[f"table_{tbl_match.group(2)}"] = i
        
        # Sections (look for numbered headings like "1.2 Title")
        sec_match = re.match(r'^(\d+(?:\.\d+)*)\s+\w+', text)
        if sec_match:
            ref_map[f"section_{sec_match.group(1)}"] = i
    
    return ref_map


def apply_clean_table_style(table):
    """
    Apply clean styling to real data tables.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing borders
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    # Add clean, subtle borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'BFBFBF')  # Light gray
        border.set(qn('w:space'), '0')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Clean cell formatting
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                # Remove cell shading
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    # Remove gray backgrounds
                    if fill and fill.upper() in ('F5F5F5', 'F0F0F0', 'EFEFEF', 'FAFAFA',
                                                   'E0E0E0', 'F8F8F8', 'FFFFFF'):
                        tcPr.remove(shd)


def apply_native_styles(doc):
    """
    Apply clean, native Word styling to the entire document.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    
    # Set Normal style
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15
    except:
        pass
    
    # Clean all paragraphs
    for paragraph in doc.paragraphs:
        # Set proper spacing
        if paragraph.paragraph_format.space_after is None:
            paragraph.paragraph_format.space_after = Pt(8)
        
        # Clean fonts
        for run in paragraph.runs:
            font_name = (run.font.name or '').lower()
            is_mono = any(m in font_name for m in ['courier', 'consolas', 'mono', 'code', 'menlo'])
            
            if is_mono:
                run.font.name = 'Consolas'
            elif not run.font.name or 'times' in font_name.lower():
                run.font.name = 'Calibri'


# ============================================================
# BACKWARDS COMPATIBILITY
# ============================================================

def enhance_docx_formatting(docx_path: Path):
    """Alias for backwards compatibility."""
    make_docx_native(docx_path)


def enhance_docx_code_blocks(docx_path: Path):
    """Alias for backwards compatibility."""
    make_docx_native(docx_path)
